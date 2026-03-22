#!/usr/bin/env python3
"""
Manuscript Review Panel — Orchestrator Script

Runs a multi-agent manuscript review pipeline using the Anthropic API.
Each agent is a Claude API call with a unique persona system prompt.
The orchestrator coordinates all agents, manages discussions, and
synthesizes the final improvement report.

Usage:
    python review_panel.py <manuscript_path> [--output <output_dir>] [--agents <agent_list>]
"""

import json
import os
import sys
import re
import time
import argparse
import subprocess
import traceback
from pathlib import Path
from datetime import datetime

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

MODEL = "claude-sonnet-4-20250514"
MAX_TOKENS = 8096

SKILL_DIR = Path(__file__).parent.parent  # manuscript-review-panel/
PERSONAS_PATH = SKILL_DIR / "references" / "personas.md"

# Agent execution order and grouping
ANTICIPATORY_AGENTS = [
    "editor_perspective",
    "strategic_advisor",
    "devils_advocate",
    "adjacent_field",
    "methods_obsessive",
    "collaborator",
    "trend_expert",
]

IMPROVEMENT_AGENTS = [
    "lab_colleague",
    "professional_editor",
    "visionary",
    "technical_expert",
    "tech_writer",
    "narrative_architect",
]

ALL_AGENTS = ANTICIPATORY_AGENTS + IMPROVEMENT_AGENTS


# ---------------------------------------------------------------------------
# Persona loader
# ---------------------------------------------------------------------------

def load_personas(personas_path: Path) -> dict:
    """Parse personas.md and return {role_label: system_prompt}."""
    content = personas_path.read_text(encoding="utf-8")
    personas = {}
    # Split on ### headings
    sections = re.split(r'\n### ', content)
    for section in sections[1:]:  # skip preamble
        # Extract role label
        label_match = re.search(r'\*\*Role label:\*\*\s*`(\w+)`', section)
        if not label_match:
            continue
        label = label_match.group(1)

        # Extract system prompt (everything between **System prompt:** and
        # the next **Structure your review as:** or end)
        prompt_match = re.search(
            r'\*\*System prompt:\*\*\s*\n(.*?)(?=\nStructure your review as:)',
            section, re.DOTALL
        )
        if not prompt_match:
            # fallback: grab everything after System prompt
            prompt_match = re.search(
                r'\*\*System prompt:\*\*\s*\n(.*)',
                section, re.DOTALL
            )
        if prompt_match:
            prompt_text = prompt_match.group(1).strip()
            # Also grab the structure instructions
            struct_match = re.search(
                r'(Structure your review as:.*)',
                section, re.DOTALL
            )
            structure = struct_match.group(1).strip() if struct_match else ""
            personas[label] = f"{prompt_text}\n\n{structure}"

    return personas


# ---------------------------------------------------------------------------
# Manuscript loader
# ---------------------------------------------------------------------------

def load_manuscript(manuscript_path: str) -> str:
    """Load manuscript text from various formats."""
    path = Path(manuscript_path)
    suffix = path.suffix.lower()

    if suffix == ".txt" or suffix == ".md":
        return path.read_text(encoding="utf-8")
    elif suffix == ".pdf":
        # Use pdftotext for text extraction
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            capture_output=True, text=True
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        # Fallback to Python
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return "\n\n".join(text_parts)
        except ImportError:
            # Try pymupdf
            import fitz
            doc = fitz.open(str(path))
            return "\n\n".join(page.get_text() for page in doc)
    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            result = subprocess.run(
                ["python3", "-c",
                 f"from docx import Document; d=Document('{path}'); print('\\n\\n'.join(p.text for p in d.paragraphs))"],
                capture_output=True, text=True
            )
            return result.stdout
    else:
        # Try reading as text
        return path.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# API caller
# ---------------------------------------------------------------------------

def call_claude(system_prompt: str, user_message: str, tools: list = None) -> str:
    """Call the Anthropic API and return the text response."""
    import anthropic
    client = anthropic.Anthropic()

    messages = [{"role": "user", "content": user_message}]

    kwargs = {
        "model": MODEL,
        "max_tokens": MAX_TOKENS,
        "system": system_prompt,
        "messages": messages,
    }

    # Add web search tool for agents that need it
    if tools:
        kwargs["tools"] = tools

    # Retry loop for API calls
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.messages.create(**kwargs)

            # Handle tool use (web search) — iterate until no more tool use
            while response.stop_reason == "tool_use":
                # Collect assistant content and make tool results
                assistant_content = response.content
                tool_results = []
                for block in assistant_content:
                    if block.type == "tool_use":
                        tool_results.append({
                            "type": "tool_result",
                            "tool_use_id": block.id,
                            "content": "Search completed. Continue with your review."
                        })

                messages.append({"role": "assistant", "content": assistant_content})
                messages.append({"role": "user", "content": tool_results})
                kwargs["messages"] = messages
                response = client.messages.create(**kwargs)

            # Extract text
            text_parts = []
            for block in response.content:
                if hasattr(block, "text"):
                    text_parts.append(block.text)
            return "\n".join(text_parts)

        except Exception as e:
            if attempt < max_retries - 1:
                wait = 2 ** attempt * 5
                print(f"  ⚠ API error (attempt {attempt+1}): {e}. Retrying in {wait}s...")
                time.sleep(wait)
            else:
                return f"[ERROR: Agent failed after {max_retries} attempts: {e}]"


# ---------------------------------------------------------------------------
# Agent runner
# ---------------------------------------------------------------------------

AGENT_DISPLAY_NAMES = {
    "editor_perspective": "📰 Journal Editor Perspective",
    "strategic_advisor": "🎯 Strategic Impact Advisor",
    "devils_advocate": "😈 Devil's Advocate (Same-Field Expert)",
    "adjacent_field": "🔭 Adjacent-Field Reviewer",
    "methods_obsessive": "🔬 Methods & Logic Specialist",
    "collaborator": "🤝 Sympathetic Collaborator",
    "trend_expert": "📊 Literature & Trends Expert",
    "narrative_architect": "🏛️ Narrative Architect (Writing Expert)",
    "lab_colleague": "🧪 Lab Colleague",
    "professional_editor": "✍️ Professional Scientific Editor",
    "visionary": "🚀 Visionary Senior Researcher",
    "technical_expert": "🛠️ Technical Methods Expert",
    "tech_writer": "📝 Technical Writer & Visualization",
}

# Agents that should have web search capability
SEARCH_ENABLED_AGENTS = {
    "devils_advocate", "trend_expert", "technical_expert",
    "methods_obsessive", "strategic_advisor", "narrative_architect",
}

WEB_SEARCH_TOOL = [{"type": "web_search_20250305", "name": "web_search"}]


def run_agent(agent_label: str, persona_prompt: str, manuscript_text: str,
              supplementary_context: str = "") -> dict:
    """Run a single agent and return its review."""
    display_name = AGENT_DISPLAY_NAMES.get(agent_label, agent_label)
    print(f"\n{'='*60}")
    print(f"  Running: {display_name}")
    print(f"{'='*60}")

    user_msg = f"""Please review the following manuscript draft. Remember: your goal
is to help the authors improve this paper for high-impact publication.
Every concern should come with a constructive suggestion.

<manuscript>
{manuscript_text}
</manuscript>
"""
    if supplementary_context:
        user_msg += f"""
<additional_context>
{supplementary_context}
</additional_context>
"""

    tools = WEB_SEARCH_TOOL if agent_label in SEARCH_ENABLED_AGENTS else None

    start = time.time()
    review_text = call_claude(persona_prompt, user_msg, tools=tools)
    elapsed = time.time() - start

    print(f"  ✓ {display_name} completed in {elapsed:.1f}s")
    print(f"  Review length: {len(review_text)} chars")

    return {
        "agent": agent_label,
        "display_name": display_name,
        "review": review_text,
        "elapsed_seconds": round(elapsed, 1),
    }


# ---------------------------------------------------------------------------
# Discussion phase
# ---------------------------------------------------------------------------

def identify_disagreements(reviews: list[dict]) -> str:
    """Ask Claude to identify disagreements among reviews."""
    reviews_text = ""
    for r in reviews:
        reviews_text += f"\n\n--- {r['display_name']} ---\n{r['review']}"

    system = """You are analyzing multiple reviews of the same manuscript to
identify points of disagreement. Focus on substantive disagreements about:
- Whether specific aspects of the methodology are adequate
- Whether certain experiments/analyses are necessary
- Whether claims are supported by the data
- How the paper should be framed or positioned
- Priority of improvements

For each disagreement, note which reviewers disagree and summarize each position.
Output as a numbered list of disagreement points. If there are no meaningful
disagreements, say so. Be concise."""

    user_msg = f"""Here are reviews from multiple perspectives on the same manuscript.
Identify the key points where reviewers disagree:

{reviews_text}"""

    return call_claude(system, user_msg)


def run_discussion(disagreements: str, reviews: list[dict]) -> str:
    """Run a moderated discussion on disagreement points."""
    reviews_summary = ""
    for r in reviews:
        reviews_summary += f"\n\n--- {r['display_name']} ---\n{r['review'][:2000]}..."

    system = """You are moderating a discussion among manuscript reviewers about
points of disagreement. For each disagreement:

1. State the disagreement clearly
2. Present the strongest argument for each side
3. Classify it as: factual (resolvable), methodological preference (both valid),
   or significance judgment (subjective)
4. Provide a clear, actionable recommendation for the authors

The goal is to give authors clear guidance even where reviewers disagree.
Keep the discussion focused and conclude each point with a concrete suggestion."""

    user_msg = f"""Points of disagreement identified:

{disagreements}

Context from the reviews:
{reviews_summary}

Please moderate a discussion on each disagreement and provide actionable
recommendations for the authors."""

    print(f"\n{'='*60}")
    print(f"  Running: 💬 Panel Discussion on Disagreements")
    print(f"{'='*60}")

    result = call_claude(system, user_msg)
    print(f"  ✓ Discussion completed")
    return result


# ---------------------------------------------------------------------------
# Final synthesis
# ---------------------------------------------------------------------------

def synthesize_final_report(reviews: list[dict], discussion: str,
                            manuscript_text: str, meta_persona: str) -> str:
    """Run the meta-reviewer to produce the final synthesis report."""
    print(f"\n{'='*60}")
    print(f"  Running: 📋 Final Report Synthesis")
    print(f"{'='*60}")

    reviews_text = ""
    for r in reviews:
        reviews_text += f"\n\n{'='*40}\n{r['display_name']}\n{'='*40}\n{r['review']}"

    user_msg = f"""You have received reviews from a panel of 12 specialist reviewers
who have all read the same manuscript. Your job is to synthesize all their
feedback into a single, comprehensive improvement roadmap for the authors.

Here is the manuscript:

<manuscript>
{manuscript_text[:5000]}...
[manuscript continues — focus on the reviews below]
</manuscript>

Here are all the individual reviews:

<reviews>
{reviews_text}
</reviews>

Here are the results of a moderated discussion on points of disagreement:

<discussion>
{discussion}
</discussion>

Now produce the final synthesized improvement report following your
structured format. Be thorough, specific, and actionable. The authors
should be able to use this report as a roadmap for revision."""

    return call_claude(meta_persona, user_msg)


# ---------------------------------------------------------------------------
# Report formatter — HTML with Mermaid support
# ---------------------------------------------------------------------------

def _md_to_html(md_text: str) -> str:
    """Convert Markdown text to HTML.

    Uses the ``markdown`` library for robust parsing (headings, bold, italic,
    code blocks, tables, nested lists, etc.).  Mermaid fenced code blocks are
    post-processed into ``<div class="mermaid">`` so the Mermaid JS library
    renders them automatically.
    """
    import markdown as _md
    import re as _re

    html = _md.markdown(
        md_text,
        extensions=["tables", "fenced_code", "toc", "nl2br"],
    )

    # Post-process: convert <pre><code class="language-mermaid">…</code></pre>
    # into <div class="mermaid">…</div> so Mermaid JS picks them up.
    html = _re.sub(
        r'<pre><code\s+class="language-mermaid">(.*?)</code></pre>',
        r'<div class="mermaid">\1</div>',
        html,
        flags=_re.DOTALL,
    )
    html = _re.sub(
        r'<pre><code\s+class="mermaid">(.*?)</code></pre>',
        r'<div class="mermaid">\1</div>',
        html,
        flags=_re.DOTALL,
    )
    return html


HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Review Panel Report — {manuscript_name}</title>
<!-- Mermaid.js for diagram rendering -->
<script src="https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.min.js"></script>
<script>mermaid.initialize({{startOnLoad: true, theme: 'neutral'}});</script>
<style>
  :root {{
    --bg: #ffffff; --fg: #1a1a2e; --accent: #0f3460; --accent2: #16213e;
    --border: #ddd; --card-bg: #f8f9fa; --code-bg: #f0f0f0;
    --green: #27ae60; --orange: #f39c12; --red: #e74c3c; --blue: #2980b9;
  }}
  @media (prefers-color-scheme: dark) {{
    :root {{
      --bg: #1a1a2e; --fg: #e0e0e0; --accent: #64b5f6; --accent2: #90caf9;
      --border: #333; --card-bg: #16213e; --code-bg: #0f3460;
      --green: #66bb6a; --orange: #ffa726; --red: #ef5350; --blue: #42a5f5;
    }}
  }}
  * {{ box-sizing: border-box; }}
  body {{
    font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
    line-height: 1.7; color: var(--fg); background: var(--bg);
    max-width: 960px; margin: 0 auto; padding: 2rem 1.5rem;
  }}
  h1 {{ color: var(--accent); border-bottom: 3px solid var(--accent); padding-bottom: 0.5rem; }}
  h2 {{ color: var(--accent2); margin-top: 2.5rem; border-bottom: 1px solid var(--border); padding-bottom: 0.3rem; }}
  h3 {{ color: var(--accent); margin-top: 2rem; }}
  h4 {{ color: var(--fg); margin-top: 1.5rem; }}
  a {{ color: var(--blue); }}
  hr {{ border: none; border-top: 2px solid var(--border); margin: 2rem 0; }}
  code {{ background: var(--code-bg); padding: 0.15em 0.4em; border-radius: 3px; font-size: 0.9em; }}
  pre {{ background: var(--code-bg); padding: 1rem; border-radius: 6px; overflow-x: auto; }}
  pre code {{ background: none; padding: 0; }}
  table {{
    width: 100%; border-collapse: collapse; margin: 1rem 0;
    font-size: 0.92em;
  }}
  th, td {{ border: 1px solid var(--border); padding: 0.5rem 0.75rem; text-align: left; }}
  th {{ background: var(--card-bg); font-weight: 600; }}
  tr:nth-child(even) td {{ background: var(--card-bg); }}
  blockquote {{
    border-left: 4px solid var(--accent); margin: 1rem 0; padding: 0.5rem 1rem;
    background: var(--card-bg); border-radius: 0 6px 6px 0;
  }}
  .meta {{ color: #888; font-size: 0.9em; margin-bottom: 2rem; }}
  .reviewer-card {{
    background: var(--card-bg); border: 1px solid var(--border);
    border-radius: 8px; padding: 1.5rem; margin: 1.5rem 0;
  }}
  .reviewer-card h3 {{ margin-top: 0; }}
  .reviewer-card .elapsed {{ font-size: 0.85em; color: #888; }}
  details {{ margin: 1rem 0; }}
  details summary {{
    cursor: pointer; font-weight: 600; padding: 0.5rem;
    background: var(--card-bg); border-radius: 4px;
  }}
  details[open] summary {{ margin-bottom: 0.5rem; }}
  .priority-critical {{ border-left: 4px solid var(--red); padding-left: 1rem; }}
  .priority-high {{ border-left: 4px solid var(--orange); padding-left: 1rem; }}
  .priority-medium {{ border-left: 4px solid var(--blue); padding-left: 1rem; }}
  .mermaid {{ margin: 1.5rem 0; text-align: center; }}
  .toc {{ background: var(--card-bg); padding: 1.5rem; border-radius: 8px; margin: 1.5rem 0; }}
  .toc ul {{ list-style: none; padding-left: 1rem; }}
  .toc li {{ margin: 0.3rem 0; }}
  .toc a {{ text-decoration: none; }}
  .toc a:hover {{ text-decoration: underline; }}
  @media print {{
    body {{ max-width: 100%; font-size: 11pt; }}
    .reviewer-card {{ break-inside: avoid; }}
    details {{ display: block; }}
    details > * {{ display: block; }}
  }}
</style>
</head>
<body>

<h1>Manuscript Review Panel Report</h1>
<div class="meta">
  <strong>Manuscript:</strong> {manuscript_name}<br>
  <strong>Date:</strong> {date}<br>
  <strong>Reviewers:</strong> {num_reviewers} specialist agents<br>
  <strong>Model:</strong> {model}
</div>

<nav class="toc">
<strong>Contents</strong>
<ul>
  <li><a href="#synthesis">Synthesised Improvement Roadmap</a></li>
  <li><a href="#individual-reviews">Individual Reviews ({num_reviewers})</a></li>
  <li><a href="#panel-discussion">Panel Discussion on Disagreements</a></li>
</ul>
</nav>

<hr>

<section id="synthesis">
{synthesis_html}
</section>

<hr>

<h2 id="individual-reviews">Full Individual Reviews</h2>
<p>Click any reviewer card to expand the full review.</p>

{reviews_html}

<hr>

<h2 id="panel-discussion">Panel Discussion on Disagreements</h2>
{discussion_html}

<hr>
<footer style="text-align:center; color:#888; font-size:0.85em; margin-top:3rem;">
  Report generated by <strong>Manuscript Review Panel</strong> skill.<br>
  This review is meant to help improve the manuscript — not to judge it.<br>
  All suggestions should be evaluated by the authors based on their expert knowledge.
</footer>

</body>
</html>
"""


def format_html_report(final_synthesis: str, individual_reviews: list[dict],
                       discussion: str, manuscript_path: str) -> str:
    """Format the complete report as a self-contained HTML document with Mermaid support."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    fname = Path(manuscript_path).name

    # Build individual review cards (collapsed by default)
    review_cards: list[str] = []
    for r in individual_reviews:
        body_html = _md_to_html(r["review"])
        card = (
            f'<details class="reviewer-card">\n'
            f'  <summary>{r["display_name"]} '
            f'<span class="elapsed">({r["elapsed_seconds"]}s)</span></summary>\n'
            f'  {body_html}\n'
            f'</details>'
        )
        review_cards.append(card)

    return HTML_TEMPLATE.format(
        manuscript_name=fname,
        date=now,
        num_reviewers=len(individual_reviews),
        model=MODEL,
        synthesis_html=_md_to_html(final_synthesis),
        reviews_html="\n\n".join(review_cards),
        discussion_html=_md_to_html(discussion),
    )


# Keep a slim Markdown version for quick reference / version control
def format_markdown_report(final_synthesis: str, individual_reviews: list[dict],
                           discussion: str, manuscript_path: str) -> str:
    """Format the complete report as a Markdown document (supplementary)."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    fname = Path(manuscript_path).name

    report = f"""# Manuscript Review Panel Report

**Manuscript:** {fname}
**Date:** {now}
**Reviewers:** {len(individual_reviews)} specialist agents
**Model:** {MODEL}

---

{final_synthesis}

---

# Full Individual Reviews

Below are the complete reviews from each panel member for reference.

"""
    for r in individual_reviews:
        report += f"""
---

## {r['display_name']}

*Completed in {r['elapsed_seconds']}s*

{r['review']}

"""

    report += f"""
---

## Panel Discussion on Disagreements

{discussion}

---

*Report generated by Manuscript Review Panel skill.*
*This review is meant to help improve the manuscript — not to judge it.*
*All suggestions should be evaluated by the authors based on their expert knowledge.*
"""
    return report


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Run manuscript review panel")
    parser.add_argument("manuscript", help="Path to manuscript file (PDF, DOCX, TXT, MD)")
    parser.add_argument("--output", "-o", default=None,
                        help="Output directory (default: same as manuscript)")
    parser.add_argument("--agents", "-a", nargs="*", default=None,
                        help="Specific agents to run (default: all)")
    parser.add_argument("--context", "-c", default=None,
                        help="Path to additional context file (cover letter, etc.)")
    parser.add_argument("--skip-discussion", action="store_true",
                        help="Skip the discussion phase")
    args = parser.parse_args()

    manuscript_path = Path(args.manuscript)
    if not manuscript_path.exists():
        print(f"Error: {manuscript_path} not found")
        sys.exit(1)

    output_dir = Path(args.output) if args.output else manuscript_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # Determine which agents to run
    agents_to_run = args.agents if args.agents else ALL_AGENTS

    # Load personas
    print("Loading personas...")
    personas = load_personas(PERSONAS_PATH)
    print(f"  Loaded {len(personas)} personas")

    # Load manuscript
    print(f"Loading manuscript: {manuscript_path}")
    manuscript_text = load_manuscript(str(manuscript_path))
    print(f"  Extracted {len(manuscript_text)} characters")

    if len(manuscript_text) < 100:
        print("Warning: Very little text extracted. Check manuscript format.")

    # Load supplementary context if provided
    supp_context = ""
    if args.context:
        supp_context = Path(args.context).read_text(encoding="utf-8")
        print(f"  Loaded supplementary context: {len(supp_context)} chars")

    # ---- Phase 1: Run all agents ----
    print(f"\n{'#'*60}")
    print(f"  PHASE 1: Individual Reviews ({len(agents_to_run)} agents)")
    print(f"{'#'*60}")

    all_reviews = []
    for agent_label in agents_to_run:
        if agent_label not in personas:
            print(f"  ⚠ Unknown agent: {agent_label}, skipping")
            continue
        review = run_agent(agent_label, personas[agent_label],
                          manuscript_text, supp_context)
        all_reviews.append(review)

        # Save individual review
        review_path = output_dir / f"review_{agent_label}.md"
        review_path.write_text(
            f"# {review['display_name']}\n\n{review['review']}",
            encoding="utf-8"
        )

    # ---- Phase 2: Identify disagreements & discuss ----
    discussion_text = ""
    if not args.skip_discussion and len(all_reviews) > 2:
        print(f"\n{'#'*60}")
        print(f"  PHASE 2: Panel Discussion")
        print(f"{'#'*60}")

        print("\nIdentifying disagreements...")
        disagreements = identify_disagreements(all_reviews)
        print(f"  Disagreements identified: {len(disagreements)} chars")

        print("\nRunning moderated discussion...")
        discussion_text = run_discussion(disagreements, all_reviews)

    # ---- Phase 3: Final synthesis ----
    print(f"\n{'#'*60}")
    print(f"  PHASE 3: Final Synthesis")
    print(f"{'#'*60}")

    meta_persona = personas.get("meta_reviewer", "You are synthesizing reviews.")
    final_synthesis = synthesize_final_report(
        all_reviews, discussion_text, manuscript_text, meta_persona
    )

    # ---- Phase 4: Format and save report ----
    print(f"\n{'#'*60}")
    print(f"  PHASE 4: Formatting Report")
    print(f"{'#'*60}")

    # Primary output: HTML with Mermaid diagram support
    report_html = format_html_report(
        final_synthesis, all_reviews, discussion_text, str(manuscript_path)
    )
    html_path = output_dir / f"review_report_{manuscript_path.stem}.html"
    html_path.write_text(report_html, encoding="utf-8")
    print(f"\n✅ HTML report saved: {html_path}")

    # Supplementary: Markdown for version control / quick grep
    report_md = format_markdown_report(
        final_synthesis, all_reviews, discussion_text, str(manuscript_path)
    )
    md_path = output_dir / f"review_report_{manuscript_path.stem}.md"
    md_path.write_text(report_md, encoding="utf-8")
    print(f"✅ Markdown report saved: {md_path}")

    report_path = html_path  # primary deliverable

    # Also save raw data as JSON
    data_path = output_dir / f"review_data_{manuscript_path.stem}.json"
    data = {
        "manuscript": str(manuscript_path),
        "timestamp": datetime.now().isoformat(),
        "model": MODEL,
        "agents_run": [r["agent"] for r in all_reviews],
        "reviews": all_reviews,
        "discussion": discussion_text,
        "synthesis": final_synthesis,
    }
    data_path.write_text(json.dumps(data, indent=2, ensure_ascii=False),
                         encoding="utf-8")
    print(f"✅ Raw data saved: {data_path}")

    return str(report_path)


if __name__ == "__main__":
    main()
